"""
scope_extractor.py
-------------------
Extracts IP addresses, CIDR ranges, and hostnames referenced anywhere in
arbitrarily-formatted documents — plain text, Markdown, CSV, logs, config
files, HTML, or (with the optional libraries below installed) Word and PDF
documents.  Intended to turn a pentest/vuln-scan scoping document (an SOW,
an email, an asset list, a network diagram's text export, ...) into a
clean, structured scope list.

How "arbitrary formatting" is actually handled
------------------------------------------------
This does not try to understand document structure.  Every supported file
type is reduced to plain text, then the same three regex-and-validate
passes run over that text regardless of where it came from — a table cell,
a paragraph, a CSV row, a code block, it's all just text by the time the
extractors see it.  That's what makes formatting irrelevant: tables,
prose, bullet lists, and key/value config files are all read the same way.

Supported file types
---------------------
  Read natively as text : .txt .md .csv .tsv .log .json .yaml/.yml .ini
                           .cfg .conf .html/.htm .xml .eml ... and any file
                           with no special handler (read as UTF-8, invalid
                           bytes replaced rather than raising).
  .docx                  : requires `pip install python-docx` (optional).
  .pdf                   : requires `pip install pypdf` (optional).
  A directory             : every file inside it (use --recursive to also
                            descend into subdirectories).
  No paths given          : reads one document's text from stdin.

A file requiring a missing optional library is skipped with a clear
message rather than aborting the whole run — see read_document_text().

Extraction approach
---------------------
IPs / CIDRs : a permissive regex finds *candidate* tokens (anything
              dotted-quad- or hex-colon-shaped); each candidate is then
              validated with the stdlib `ipaddress` module, which is what
              actually rejects things like version strings ("3.10.2") or
              malformed addresses.  Regex finds candidates, ipaddress is
              the source of truth.  IPv6 support is best-effort: embedded
              IPv4-in-IPv6 notation (e.g. "::ffff:192.0.2.1") is not
              matched.

              Routing table output is explicitly handled: Palo Alto
              ("show routing route") and Juniper ("show route" / "set
              routing-options static route") already use CIDR slash
              notation throughout, so the candidate regex above is enough.
              Classic Cisco/ASA syntax instead writes a range as two
              adjacent tokens -- "<network> <netmask>" in static routes
              and "show route" output (e.g. "10.0.0.0 255.0.0.0"), or
              "<network> <wildcard-mask>" in OSPF/EIGRP "network"
              statements and ACLs (e.g. "10.0.0.0 0.255.255.255").  Both
              forms are recognised and converted to the equivalent CIDR by
              _consume_network_mask_pairs() before the plain candidate
              pass runs, so the mask token itself never leaks into the
              output as a bogus standalone "IP".

Hostnames   : a label-based regex requires the final label ("TLD") to be
              2+ purely-alphabetic characters not in a file-extension
              blocklist (see DEFAULT_EXCLUDED_TLDS).  This single rule is
              what keeps IPv4 addresses (numeric final octet), decimal
              numbers, and most filenames (audit.py, README.md, ...) from
              being misclassified as hostnames.  It is a heuristic, not a
              real public-suffix-list lookup — uncommon file extensions or
              deliberately hostname-shaped strings can still slip through.
              Use --exclude-tld to extend the blocklist for your corpus.

Usage:
    python scope_extractor.py notes.txt SOW.docx network_diagram_export.pdf
    python scope_extractor.py ./scoping_docs --recursive --out scope.yaml
    type notes.txt | python scope_extractor.py --out scope.json

Output:
    A YAML (default) or JSON file with "sources", "cidrs", "ips", and
    "hostnames" keys — format is inferred from --out's extension, or set
    explicitly with --format.  Stable, sorted output so re-running against
    an updated document set produces a clean git diff.

Extending:
    New file format    – add an entry to _READERS mapping a suffix to a
                          `Callable[[Path], str]` that returns extracted text.
    New scope category – add an extract_*() function alongside
                          extract_ip_tokens()/extract_hostnames(), a field
                          on ScopeFindings, and a --no-<category> CLI flag.
"""

import argparse
import ipaddress
import json
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable, Iterator

import yaml

# ── Errors ───────────────────────────────────────────────────────────────────

class MissingDependencyError(Exception):
    """Raised when a file format's reader needs an optional library that
    is not installed.  Caught per-file so one missing library does not
    abort scanning every other document."""


# ── Document readers ─────────────────────────────────────────────────────────

def _read_plain_text(path: Path) -> str:
    """Read any text-ish file.  Invalid bytes are replaced, never raised on,
    since the goal is to find IP/hostname substrings, not round-trip the file."""
    return path.read_text(encoding="utf-8", errors="replace")


def _read_docx(path: Path) -> str:
    """Read paragraph and table cell text from a Word document."""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise MissingDependencyError(
            f"reading .docx files requires the 'python-docx' package "
            f"(pip install python-docx)"
        ) from exc

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    """Read text from every page of a PDF."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise MissingDependencyError(
            f"reading .pdf files requires the 'pypdf' package "
            f"(pip install pypdf)"
        ) from exc

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


# Suffixes that need a dedicated reader.  Anything not listed here falls
# back to _read_plain_text — which is exactly right for the large set of
# formats that are already text under the hood (.txt, .md, .csv, .json,
# .yaml, .ini, .log, .html, .xml, .eml, ...).
_READERS: dict[str, Callable[[Path], str]] = {
    ".docx": _read_docx,
    ".pdf":  _read_pdf,
}


def read_document_text(path: Path) -> str:
    """Dispatch to the right reader for path's suffix and return its text.

    Raises:
        MissingDependencyError if the format needs an optional library
            that isn't installed.
        OSError on file read failures (permissions, etc.).
    """
    reader = _READERS.get(path.suffix.lower(), _read_plain_text)
    return reader(path)


def iter_input_files(paths: list[Path], recursive: bool) -> Iterator[Path]:
    """Expand a list of file/directory paths into individual file paths.

    A path that doesn't exist is skipped with a warning rather than
    raising, so one bad path doesn't abort scanning the rest.
    """
    for p in paths:
        if not p.exists():
            print(f"  [err] {p}: path not found", file=sys.stderr)
            continue
        if p.is_dir():
            pattern = "**/*" if recursive else "*"
            for child in sorted(p.glob(pattern)):
                if child.is_file():
                    yield child
        else:
            yield p


# ── IP / CIDR extraction ──────────────────────────────────────────────────────

# Loose candidate patterns — deliberately permissive.  ipaddress.ip_address()/
# ip_network() below is what actually decides validity, so there's no need
# (or attempt) to encode "0-255 per octet" or full IPv6 grammar in regex.
_IPV4_CANDIDATE_RE = re.compile(r"\b\d{1,3}(?:\.\d{1,3}){3}(?:/\d{1,2})?\b")
# Boundary assertions exclude word chars and ":" (things that could extend
# a real hex-group-or-colon run) but deliberately NOT ".": the consuming
# pattern never contains a literal "." anyway, so excluding it from the
# boundary check buys no protection -- it only ever caused a real bug,
# where a CIDR range immediately followed by sentence punctuation (e.g.
# "...covers 2001:db8::/32." at the end of a sentence) had its "/32"
# silently dropped because the lookahead failed right after the prefix
# length and the optional CIDR group backtracked away.
_IPV6_CANDIDATE_RE = re.compile(
    r"(?<![\w:])(?:[0-9A-Fa-f]{0,4}:){2,7}[0-9A-Fa-f]{0,4}(?:/\d{1,3})?(?![\w:])"
)

# Classic Cisco/ASA routing syntax expresses a range as two adjacent IPv4
# tokens — "<network> <netmask>" in static routes and "show route"/"show
# running-config" output (e.g. "10.0.0.0 255.0.0.0"), or "<network>
# <wildcard-mask>" (the bitwise inverse of a netmask) in OSPF/EIGRP
# "network" statements and ACLs (e.g. "10.0.0.0 0.255.255.255").  Palo Alto
# and Juniper output, by contrast, already uses CIDR slash notation
# end-to-end, so it's handled by the plain candidate regexes above with no
# extra work.  Build {dotted-quad mask string: prefix length} lookups for
# every /0-/32 netmask and its wildcard counterpart so both forms convert
# to the equivalent CIDR.
def _build_netmask_tables() -> tuple[dict[str, int], dict[str, int]]:
    netmasks: dict[str, int] = {}
    wildcards: dict[str, int] = {}
    for prefix in range(0, 33):
        net = ipaddress.IPv4Network(f"0.0.0.0/{prefix}")
        netmasks[str(net.netmask)] = prefix
        wildcards[str(net.hostmask)] = prefix
    return netmasks, wildcards


_NETMASK_TO_PREFIX, _WILDCARD_TO_PREFIX = _build_netmask_tables()

_NETWORK_MASK_PAIR_RE = re.compile(
    r"\b(\d{1,3}(?:\.\d{1,3}){3})[ \t]+(\d{1,3}(?:\.\d{1,3}){3})\b"
)


def _consume_network_mask_pairs(text: str) -> tuple[set[str], str]:
    """Find "<network> <mask>" pairs and convert each to CIDR.

    Only pairs whose second token is an exact /0-/32 netmask or wildcard
    mask are converted (e.g. a coincidental "192.168.1.1 192.168.1.2"
    elsewhere in the text is untouched — neither token is a valid mask, so
    the regex match is returned as-is).  Matched pairs are blanked out
    (replaced with equal-length whitespace, so positions/line numbers
    elsewhere are unaffected) so the mask token doesn't also get picked up
    as a bogus standalone "IP" by the plain candidate pass that follows.

    Only matches within a single line (uses [ \\t]+, not \\s+) so a mask at
    the end of one line is never paired with an unrelated address that
    happens to start the next line.

    Returns:
        (cidrs found, text with matched pairs blanked out).
    """
    cidrs: set[str] = set()

    def _replace(m: re.Match) -> str:
        network_str, mask_str = m.group(1), m.group(2)
        prefix = _NETMASK_TO_PREFIX.get(mask_str)
        if prefix is None:
            prefix = _WILDCARD_TO_PREFIX.get(mask_str)
        if prefix is None:
            return m.group(0)  # second token isn't a real mask — leave as-is

        try:
            network = ipaddress.ip_network(f"{network_str}/{prefix}", strict=False)
        except ValueError:
            return m.group(0)

        cidrs.add(str(network))
        return " " * len(m.group(0))

    return cidrs, _NETWORK_MASK_PAIR_RE.sub(_replace, text)


def extract_ip_tokens(text: str) -> tuple[set[str], set[str]]:
    """Find IPv4/IPv6 addresses and CIDR ranges referenced in text.

    Returns:
        (ips, cidrs) — bare addresses and prefixed ranges, kept separate so
        callers/output can list them under distinct headings.
    """
    ips: set[str] = set()

    cidrs, text = _consume_network_mask_pairs(text)

    candidates = _IPV4_CANDIDATE_RE.findall(text) + _IPV6_CANDIDATE_RE.findall(text)
    for candidate in candidates:
        try:
            if "/" in candidate:
                network = ipaddress.ip_network(candidate, strict=False)
                cidrs.add(str(network))
            else:
                ips.add(str(ipaddress.ip_address(candidate)))
        except ValueError:
            continue  # not actually a valid address/network — discard

    return ips, cidrs


# ── Hostname extraction ───────────────────────────────────────────────────────

# Common file extensions that would otherwise pass the "2+ alpha-only
# label" TLD test below and get misclassified as hostnames (audit.py,
# README.md, requirements.txt, ...).  Extend with --exclude-tld.
DEFAULT_EXCLUDED_TLDS: set[str] = {
    "py", "txt", "md", "json", "yaml", "yml", "csv", "tsv", "log", "xml",
    "docx", "doc", "pdf", "xlsx", "xls", "pptx", "ppt",
    "png", "jpg", "jpeg", "gif", "bmp", "svg", "ico",
    "zip", "tar", "gz", "bz2", "rar", "exe", "dll", "so",
    "ini", "cfg", "conf", "env", "sh", "bat", "cmd", "ps1",
    "html", "htm", "css", "js", "ts", "jar", "class", "sql",
    "db", "bak", "tmp", "dat",
}

_HOSTNAME_LABEL = r"[A-Za-z0-9](?:[A-Za-z0-9-]{0,61}[A-Za-z0-9])?"
_HOSTNAME_RE = re.compile(rf"\b(?:{_HOSTNAME_LABEL}\.)+{_HOSTNAME_LABEL}\b")


def extract_hostnames(text: str, excluded_tlds: set[str]) -> set[str]:
    """Find hostname-like tokens in text.

    See the module docstring's "Hostnames" section for the TLD-based
    heuristic this relies on.
    """
    hostnames: set[str] = set()
    for match in _HOSTNAME_RE.findall(text):
        tld = match.rsplit(".", 1)[-1]
        if len(tld) < 2 or not tld.isalpha():
            continue
        if tld.lower() in excluded_tlds:
            continue
        hostnames.add(match.rstrip("."))
    return hostnames


# ── Aggregation ────────────────────────────────────────────────────────────────

@dataclass
class ScopeFindings:
    """Accumulated extraction results, mergeable across multiple documents."""
    ips:       set[str] = field(default_factory=set)
    cidrs:     set[str] = field(default_factory=set)
    hostnames: set[str] = field(default_factory=set)

    def update(self, other: "ScopeFindings") -> None:
        self.ips |= other.ips
        self.cidrs |= other.cidrs
        self.hostnames |= other.hostnames


def extract_from_text(
    text: str,
    excluded_tlds: set[str],
    include_ips: bool = True,
    include_cidrs: bool = True,
    include_hostnames: bool = True,
) -> ScopeFindings:
    """Run the requested extractors over one document's text."""
    ips: set[str] = set()
    cidrs: set[str] = set()
    if include_ips or include_cidrs:
        found_ips, found_cidrs = extract_ip_tokens(text)
        ips = found_ips if include_ips else set()
        cidrs = found_cidrs if include_cidrs else set()

    hostnames = extract_hostnames(text, excluded_tlds) if include_hostnames else set()

    return ScopeFindings(ips=ips, cidrs=cidrs, hostnames=hostnames)


# ── Sorting helpers ───────────────────────────────────────────────────────────

def _ip_sort_key(ip_str: str) -> tuple:
    ip = ipaddress.ip_address(ip_str)
    return (ip.version, ip)


def _network_sort_key(net_str: str) -> tuple:
    net = ipaddress.ip_network(net_str)
    return (net.version, net.network_address, net.prefixlen)


# ── Output ────────────────────────────────────────────────────────────────────

def build_payload(findings: ScopeFindings, sources: list[str]) -> dict:
    """Assemble the final structured output, sorted for stable, diffable runs."""
    return {
        "generated":  datetime.now(timezone.utc).isoformat(),
        "sources":    sources,
        "cidrs":      sorted(findings.cidrs, key=_network_sort_key),
        "ips":        sorted(findings.ips, key=_ip_sort_key),
        "hostnames":  sorted(findings.hostnames, key=str.lower),
    }


def write_payload(payload: dict, dest: Path, fmt: str) -> None:
    if fmt == "json":
        dest.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    else:
        dest.write_text(
            yaml.dump(payload, sort_keys=False, default_flow_style=False, allow_unicode=True),
            encoding="utf-8",
        )


# ── Entry point ───────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Extract IP addresses, CIDR ranges, and hostnames referenced in "
            "arbitrarily-formatted documents."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "examples:\n"
            "  %(prog)s notes.txt SOW.docx\n"
            "  %(prog)s ./scoping_docs --recursive --out scope.yaml\n"
            "  type notes.txt | %(prog)s --out scope.json\n"
        ),
    )
    parser.add_argument(
        "paths", nargs="*",
        help="Files or directories to scan. Omit to read one document from stdin.",
    )
    parser.add_argument(
        "--recursive", "-r", action="store_true",
        help="When a path is a directory, also descend into subdirectories.",
    )
    parser.add_argument(
        "--out", "-o", default="scope.yaml",
        help="Output path (default: scope.yaml).",
    )
    parser.add_argument(
        "--format", "-f", choices=["yaml", "json"], default=None,
        help="Output format. Default: inferred from --out's extension (.json -> json, else yaml).",
    )
    parser.add_argument(
        "--exclude-tld", action="append", default=[], metavar="EXT",
        help=(
            "Additional file-extension-like suffix to exclude from hostname "
            "matches (e.g. --exclude-tld rst). Repeatable."
        ),
    )
    parser.add_argument("--no-ips", action="store_true", help="Skip IP address extraction.")
    parser.add_argument("--no-cidrs", action="store_true", help="Skip CIDR range extraction.")
    parser.add_argument("--no-hostnames", action="store_true", help="Skip hostname extraction.")
    args = parser.parse_args()

    excluded_tlds = DEFAULT_EXCLUDED_TLDS | {e.lower().lstrip(".") for e in args.exclude_tld}

    findings = ScopeFindings()
    sources: list[str] = []

    if not args.paths:
        print("Reading from stdin ...")
        text = sys.stdin.read()
        findings.update(extract_from_text(
            text, excluded_tlds,
            include_ips=not args.no_ips,
            include_cidrs=not args.no_cidrs,
            include_hostnames=not args.no_hostnames,
        ))
        sources.append("<stdin>")
    else:
        for file_path in iter_input_files([Path(p) for p in args.paths], recursive=args.recursive):
            try:
                text = read_document_text(file_path)
            except (OSError, MissingDependencyError) as exc:
                print(f"  [err] {file_path}: {exc}", file=sys.stderr)
                continue

            findings.update(extract_from_text(
                text, excluded_tlds,
                include_ips=not args.no_ips,
                include_cidrs=not args.no_cidrs,
                include_hostnames=not args.no_hostnames,
            ))
            sources.append(str(file_path))
            print(f"  [ok]  {file_path}")

    if not sources:
        sys.exit("Error: no documents were successfully read — nothing to extract.")

    payload = build_payload(findings, sources)
    out_path = Path(args.out)
    fmt = args.format or ("json" if out_path.suffix.lower() == ".json" else "yaml")

    try:
        write_payload(payload, out_path, fmt)
    except OSError as exc:
        sys.exit(f"Error: failed to write {out_path}: {exc}")

    print(f"\n{len(sources)} document(s) scanned")
    print(f"  ips       : {len(findings.ips)}")
    print(f"  cidrs     : {len(findings.cidrs)}")
    print(f"  hostnames : {len(findings.hostnames)}")
    print(f"\nScope written to {out_path}")


if __name__ == "__main__":
    main()
